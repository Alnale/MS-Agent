"""DocReader — lightweight document text extraction service.

Provides a single POST /read endpoint that extracts text from
PDF, DOCX, and DOC files and returns structured JSON.

Dependencies: flask, pymupdf (fitz), python-docx
Optional:     textract (for legacy .doc files)
"""

import os
import sys
import json
import tempfile
import traceback
from pathlib import Path

from flask import Flask, request, jsonify

app = Flask(__name__)

# ---------------------------------------------------------------------------
# Text extraction backends
# ---------------------------------------------------------------------------

def read_pdf(path: str, pages: str | None = None) -> dict:
    """Extract text, tables, images, metadata, links, bookmarks, and annotations from a PDF using PyMuPDF."""
    import fitz  # pymupdf

    doc = fitz.open(path)
    total = doc.page_count

    # Document metadata
    metadata = {}
    if doc.metadata:
        for key in ("title", "author", "subject", "creator", "producer", "creationDate", "modDate"):
            val = doc.metadata.get(key)
            if val and str(val).strip():
                metadata[key] = str(val).strip()

    # Extract bookmarks/table of contents
    toc = doc.get_toc()
    bookmarks = []
    for level, title, page_num in toc:
        bookmarks.append({
            "level": level,
            "title": title,
            "page": page_num
        })

    # Parse page range: "1-3,5,7-" → [0,1,2,4,6,...]  (1-indexed input)
    page_indices = _parse_page_range(pages, total)

    sections = []
    total_images = 0
    total_tables = 0
    total_links = 0
    total_annotations = 0
    all_links = []
    all_annotations = []

    for idx in page_indices:
        page = doc[idx]
        page_parts = []

        # Count images on this page
        img_list = page.get_images()
        page_img_count = len(img_list)
        total_images += page_img_count

        # Extract tables using PyMuPDF's table finder
        try:
            tables = page.find_tables()
            if tables and tables.tables:
                total_tables += len(tables.tables)
                for t_idx, table in enumerate(tables.tables):
                    page_parts.append(f"--- 表格 {t_idx + 1} ---")
                    data = table.extract()
                    for row_idx, row in enumerate(data):
                        cells = [str(c).strip() if c else "" for c in row]
                        if any(cells):
                            page_parts.append(" | ".join(cells))
        except Exception:
            pass  # Fall back to plain text only

        # Extract text blocks for better structure preservation
        try:
            blocks = page.get_text("blocks")
            for block in blocks:
                if block[6] == 0:  # text block (not image)
                    txt = block[4].strip()
                    if txt:
                        page_parts.append(txt)
        except Exception:
            text = page.get_text("text").strip()
            if text:
                page_parts.append(text)

        # Extract hyperlinks
        try:
            links = page.get_links()
            for link in links:
                if link.get("kind") == fitz.LINK_URI:  # External URL
                    total_links += 1
                    link_info = {
                        "type": "uri",
                        "url": link.get("uri", ""),
                        "page": idx + 1,
                        "rect": link.get("rect")
                    }
                    all_links.append(link_info)
                elif link.get("kind") == fitz.LINK_GOTO:  # Internal link
                    total_links += 1
                    link_info = {
                        "type": "goto",
                        "page": link.get("page", 0) + 1,
                        "page": idx + 1,
                        "rect": link.get("rect")
                    }
                    all_links.append(link_info)
        except Exception:
            pass

        # Extract annotations
        try:
            annots = page.annots()
            if annots:
                for annot in annots:
                    total_annotations += 1
                    annot_info = {
                        "type": annot.type[1] if annot.type else "unknown",
                        "content": annot.info.get("content", ""),
                        "page": idx + 1,
                        "rect": annot.rect
                    }
                    all_annotations.append(annot_info)
        except Exception:
            pass

        if page_img_count > 0:
            page_parts.append(f"[本页包含 {page_img_count} 张图片]")

        if page_parts:
            sections.append({"page": idx + 1, "text": "\n\n".join(page_parts)})

    doc.close()

    result = {
        "format": "pdf",
        "total_pages": total,
        "extracted_pages": len(sections),
        "sections": sections,
        "text": "\n\n".join(s["text"] for s in sections),
        "images": total_images,
        "tables": total_tables,
        "links": total_links,
        "annotations": total_annotations,
    }

    # Add bookmarks if available
    if bookmarks:
        result["bookmarks"] = bookmarks

    # Add detailed links if available
    if all_links:
        # Remove rect info for cleaner output
        for link in all_links:
            link.pop("rect", None)
        result["link_details"] = all_links

    # Add detailed annotations if available
    if all_annotations:
        # Remove rect info for cleaner output
        for annot in all_annotations:
            annot.pop("rect", None)
        result["annotation_details"] = all_annotations

    if metadata:
        result["metadata"] = metadata
    return result


def read_docx(path: str) -> dict:
    """Extract text, headings, lists, tables, images, metadata, links, footnotes, endnotes, and comments from a DOCX."""
    from docx import Document
    from docx.oxml.ns import qn
    import zipfile
    import xml.etree.ElementTree as ET

    doc = Document(path)

    # Document core properties
    metadata = {}
    cp = doc.core_properties
    for attr, key in [
        ("title", "title"), ("author", "author"), ("subject", "subject"),
        ("created", "created"), ("modified", "modified"), ("revision", "revision"),
    ]:
        val = getattr(cp, attr, None)
        if val is not None:
            metadata[key] = str(val)

    # Count embedded images (inline shapes + floating shapes)
    image_count = 0
    try:
        image_count = len(doc.inline_shapes)
    except Exception:
        pass
    # Also count images in the XML that inline_shapes might miss
    try:
        blips = doc.element.findall('.//' + qn('a:blip'))
        image_count = max(image_count, len(blips))
    except Exception:
        pass

    # Extract hyperlinks from document.xml.rels
    hyperlinks = []
    try:
        with zipfile.ZipFile(path, 'r') as z:
            if 'word/_rels/document.xml.rels' in z.namelist():
                rels_xml = z.read('word/_rels/document.xml.rels')
                rels_root = ET.fromstring(rels_xml)
                ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                for rel in rels_root.findall('.//r:Relationship', ns):
                    rel_type = rel.get('Type', '')
                    if 'hyperlink' in rel_type.lower():
                        hyperlinks.append({
                            "id": rel.get('Id'),
                            "target": rel.get('Target')
                        })
    except Exception:
        pass

    # Extract footnotes and endnotes
    footnotes = []
    endnotes = []
    try:
        with zipfile.ZipFile(path, 'r') as z:
            # Footnotes
            if 'word/footnotes.xml' in z.namelist():
                fn_xml = z.read('word/footnotes.xml')
                fn_root = ET.fromstring(fn_xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for footnote in fn_root.findall('.//w:footnote', ns):
                    fn_id = footnote.get(qn('w:id'))
                    fn_text = ""
                    for p in footnote.findall('.//w:p', ns):
                        for r in p.findall('.//w:r', ns):
                            for t in r.findall('.//w:t', ns):
                                if t.text:
                                    fn_text += t.text
                    if fn_text.strip():
                        footnotes.append({"id": fn_id, "text": fn_text.strip()})

            # Endnotes
            if 'word/endnotes.xml' in z.namelist():
                en_xml = z.read('word/endnotes.xml')
                en_root = ET.fromstring(en_xml)
                for endnote in en_root.findall('.//w:endnote', ns):
                    en_id = endnote.get(qn('w:id'))
                    en_text = ""
                    for p in endnote.findall('.//w:p', ns):
                        for r in p.findall('.//w:r', ns):
                            for t in r.findall('.//w:t', ns):
                                if t.text:
                                    en_text += t.text
                    if en_text.strip():
                        endnotes.append({"id": en_id, "text": en_text.strip()})
    except Exception:
        pass

    # Extract comments
    comments = []
    try:
        with zipfile.ZipFile(path, 'r') as z:
            if 'word/comments.xml' in z.namelist():
                cm_xml = z.read('word/comments.xml')
                cm_root = ET.fromstring(cm_xml)
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for comment in cm_root.findall('.//w:comment', ns):
                    cm_id = comment.get(qn('w:id'))
                    cm_author = comment.get(qn('w:author'), "")
                    cm_date = comment.get(qn('w:date'), "")
                    cm_text = ""
                    for p in comment.findall('.//w:p', ns):
                        for r in p.findall('.//w:r', ns):
                            for t in r.findall('.//w:t', ns):
                                if t.text:
                                    cm_text += t.text
                    if cm_text.strip():
                        comments.append({
                            "id": cm_id,
                            "author": cm_author,
                            "date": cm_date,
                            "text": cm_text.strip()
                        })
    except Exception:
        pass

    # Build structured text with heading hierarchy and list markers
    lines = []
    paragraph_count = 0
    heading_count = 0
    list_count = 0
    bold_count = 0
    italic_count = 0
    underline_count = 0

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue

        paragraph_count += 1
        style_name = (para.style.name or "").lower()
        prefix = ""

        # Detect headings: "Heading 1", "标题 1", etc.
        if "heading" in style_name or style_name.startswith("标题"):
            # Extract heading level number
            level = 1
            for ch in style_name:
                if ch.isdigit():
                    level = int(ch)
                    break
            level = min(level, 6)
            prefix = "#" * level + " "
            heading_count += 1
        # Detect list items
        elif "list" in style_name or "bullet" in style_name:
            # Check indent level for sub-lists
            indent = 0
            try:
                indent_val = para.paragraph_format.left_indent
                if indent_val:
                    indent = min(int(indent_val / 360000), 4)  # EMU to rough level
            except Exception:
                pass
            prefix = "  " * indent + "- "
            list_count += 1
        # Detect numbering (ordered lists) via XML
        elif para._element.find(qn('w:numPr')) is not None:
            indent = 0
            try:
                indent_val = para.paragraph_format.left_indent
                if indent_val:
                    indent = min(int(indent_val / 360000), 4)
            except Exception:
                pass
            prefix = "  " * indent + "• "
            list_count += 1

        # Extract text formatting (bold, italic, underline)
        try:
            for run in para.runs:
                if run.bold:
                    bold_count += 1
                if run.italic:
                    italic_count += 1
                if run.underline:
                    underline_count += 1
        except Exception:
            pass

        lines.append(prefix + txt)

    # Extract tables with structure
    table_texts = []
    table_count = len(doc.tables)
    for t_idx, table in enumerate(doc.tables):
        table_texts.append(f"--- 表格 {t_idx + 1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    all_text = "\n\n".join(lines)
    if table_texts:
        all_text += "\n\n" + "\n".join(table_texts)
    if image_count > 0:
        all_text += f"\n\n[文档包含 {image_count} 张图片]"

    result = {
        "format": "docx",
        "paragraphs": paragraph_count,
        "headings": heading_count,
        "lists": list_count,
        "tables": table_count,
        "images": image_count,
        "text": all_text,
        "metadata": metadata,
        "bold_count": bold_count,
        "italic_count": italic_count,
        "underline_count": underline_count,
    }

    # Add hyperlinks if available
    if hyperlinks:
        result["hyperlinks"] = hyperlinks

    # Add footnotes if available
    if footnotes:
        result["footnotes"] = footnotes

    # Add endnotes if available
    if endnotes:
        result["endnotes"] = endnotes

    # Add comments if available
    if comments:
        result["comments"] = comments

    return result


def read_doc(path: str) -> dict:
    """Extract text from a legacy .doc file.

    Tries antiword first (fast, no Python deps), then falls back to textract.
    If neither is available, returns an error with guidance.
    """
    # 1. Try antiword (common on Linux, rare on Windows)
    try:
        import subprocess
        result = subprocess.run(
            ["antiword", path],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return {"format": "doc", "text": result.stdout.strip()}
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # 2. Try textract (pip install textract)
    try:
        import textract
        raw = textract.process(path)
        text = raw.decode("utf-8", errors="replace").strip()
        if text:
            return {"format": "doc", "text": text}
    except ImportError:
        pass
    except Exception:
        pass

    # 3. Try python-docx (some .doc files are actually docx)
    try:
        from docx import Document
        doc = Document(path)
        text = "\n\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if text:
            return {"format": "doc", "text": text}
    except Exception:
        pass

    return {
        "error": "无法读取 .doc 文件。请安装 antiword 或将文件转换为 DOCX/PDF 后重试。",
        "hint": "推荐: 使用 docflow 工具先将 .doc 转换为 .docx 或 .pdf",
    }


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _parse_page_range(spec: str | None, total: int) -> list[int]:
    """Parse a 1-indexed page range string into 0-indexed page indices.

    Examples: "1-3,5" → [0,1,2,4], "3-" → [2,3,...,total-1], None → all pages
    """
    if not spec or not spec.strip():
        return list(range(total))

    indices = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            lo, hi = part.split("-", 1)
            lo = max(int(lo), 1) if lo.strip() else 1
            hi = min(int(hi), total) if hi.strip() else total
            for i in range(lo, hi + 1):
                if 1 <= i <= total:
                    indices.add(i - 1)
        else:
            i = int(part)
            if 1 <= i <= total:
                indices.add(i - 1)
    return sorted(indices)


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/health")
def health():
    return jsonify({"status": "ok", "service": "docreader"})


@app.route("/read", methods=["POST"])
def read_document():
    """Read and extract text from a document.

    JSON body:
        path  (required): absolute or relative file path
        pages (optional): page range for PDF, e.g. "1-3,5,7-"
    """
    data = request.get_json(silent=True) or {}
    file_path = data.get("path", "").strip()
    pages = data.get("pages")

    if not file_path:
        return jsonify({"error": "缺少 path 参数"}), 400

    # Resolve relative paths
    if not os.path.isabs(file_path):
        file_path = os.path.abspath(file_path)

    if not os.path.isfile(file_path):
        return jsonify({"error": f"文件不存在: {file_path}"}), 404

    ext = Path(file_path).suffix.lower()

    try:
        if ext == ".pdf":
            result = read_pdf(file_path, pages)
        elif ext == ".docx":
            result = read_docx(file_path)
        elif ext == ".doc":
            result = read_doc(file_path)
        else:
            return jsonify({"error": f"不支持的文件格式: {ext}。支持 PDF、DOCX、DOC。"}), 400

        # Add file metadata
        result["file"] = os.path.basename(file_path)
        result["size"] = os.path.getsize(file_path)
        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": f"读取失败: {str(e)}"}), 500


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    port = int(os.environ.get("DOCREADER_PORT", "5002"))
    print(f"DocReader service starting on port {port}...")
    app.run(host="127.0.0.1", port=port, debug=False)
